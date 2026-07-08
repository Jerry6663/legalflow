"""Document parsing service — PDF and DOCX."""


from pathlib import Path
import pdfplumber
from docx import Document
from PIL import Image


class DocumentParsingError(Exception):
    pass


class DocumentParser:
    """Parse legal documents (contracts) from various formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png"}

    def __init__(self):
        pass

    def parse(self, file_path: str) -> tuple[str, dict]:
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        parsers = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,
            ".jpg": self._parse_image,
            ".jpeg": self._parse_image,
            ".png": self._parse_image,
        }

        parser = parsers.get(ext)
        if not parser:
            raise DocumentParsingError(f"No parser found for {ext}")

        try:
            return parser(file_path)
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse {file_path}: {e}")

    def _parse_pdf(self, file_path: str) -> tuple[str, dict]:
        """Parse PDF using pdfplumber."""
        with pdfplumber.open(file_path) as pdf:
            metadata = {
                "pages": len(pdf.pages),
                "format": "pdf",
            }
            text_parts = []
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        text += "\n[表格内容]\n"
                        for row in table:
                            text += " | ".join(cell or "" for cell in row) + "\n"
                if text.strip():
                    text_parts.append(f"--- 第{i}页 ---\n{text}")
        return "\n".join(text_parts), metadata

    def _parse_docx(self, file_path: str) -> tuple[str, dict]:
        """Parse Word document."""
        doc = Document(file_path)
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "format": "docx",
        }
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for i, table in enumerate(doc.tables, 1):
            text_parts.append(f"\n[表格{i}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text_parts.append(" | ".join(cells))
        return "\n".join(text_parts), metadata

    def _parse_image(self, file_path: str) -> tuple[str, dict]:
        """Basic image info — full OCR requires tesseract installation."""
        img = Image.open(file_path)
        metadata = {
            "format": "image",
            "size": img.size,
            "mode": img.mode,
            "warning": "OCR requires tesseract installation on server",
        }
        return f"[Image: {img.size[0]}x{img.size[1]}, {img.mode}]", metadata

    def parse_contract(self, file_path: str) -> dict:
        text, metadata = self.parse(file_path)
        sections = self._detect_sections(text)
        return {
            "raw_text": text,
            "metadata": metadata,
            "sections": sections,
            "word_count": len(text),
            "char_count": len(text),
        }

    def _detect_sections(self, text: str) -> list[dict]:
        section_keywords = [
            "第一条", "第二条", "第三条", "第四条", "第五条",
            "第1条", "第2条", "第3条",
            "一、", "二、", "三、", "四、", "五、",
            "1.", "2.", "3.", "4.", "5.",
            "甲方", "乙方",
        ]
        lines = text.split("\n")
        sections = []
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            for kw in section_keywords:
                if kw in line and len(line) < 100:
                    sections.append({"line_number": i + 1, "title": line[:80]})
                    break
        return sections


document_parser = DocumentParser()
